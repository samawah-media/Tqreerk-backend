"""Generate `Taqreerk_RAG_Architecture.docx` describing the RAG pipeline.

Run from repo root:
    python docs/generate_rag_doc.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


OUT = Path(__file__).resolve().parent / "Taqreerk_RAG_Architecture_v2.docx"


# ── styling helpers ────────────────────────────────────────────────────

PRIMARY = RGBColor(0x1F, 0x3A, 0x5F)   # deep navy
ACCENT = RGBColor(0x2E, 0x86, 0xAB)    # teal
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_FILL = "EAF2F8"
DARK_HDR = "1F3A5F"


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = PRIMARY
        run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    if p.runs:
        p.runs[0].text = ""
    if ":" in text and len(text.split(":", 1)[0]) <= 32:
        label, body = text.split(":", 1)
        r1 = p.add_run(label + ":")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = "Calibri"
        r2 = p.add_run(body)
        r2.font.size = Pt(11)
        r2.font.name = "Calibri"
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F8")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", "60"), ("bottom", "60"), ("left", "120"), ("right", "120")):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), val)
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def add_two_col_table(doc: Document, header: tuple[str, str], rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    h0, h1 = table.rows[0].cells
    for cell, txt in ((h0, header[0]), (h1, header[1])):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        shade_cell(cell, DARK_HDR)
    for i, (left, right) in enumerate(rows, start=1):
        c0, c1 = table.rows[i].cells
        c0.text = left
        c1.text = right
        for cell in (c0, c1):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = "Calibri"


def add_three_col_table(doc: Document, headers: tuple[str, str, str], rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Light Grid Accent 1"
    for cell, txt in zip(table.rows[0].cells, headers):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        shade_cell(cell, DARK_HDR)
    for i, row_vals in enumerate(rows, start=1):
        for cell, val in zip(table.rows[i].cells, row_vals):
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"


def add_arrow_flow(doc: Document, steps: list[str]) -> None:
    cells_per_row = 2 * len(steps) - 1
    table = doc.add_table(rows=1, cols=cells_per_row)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, step in enumerate(steps):
        cell = table.rows[0].cells[2 * i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(step)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = PRIMARY
        run.font.name = "Calibri"
        shade_cell(cell, LIGHT_FILL)
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            node = OxmlElement(f"w:{side}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "8")
            node.set(qn("w:color"), "2E86AB")
            borders.append(node)
        tc_pr.append(borders)
        if i < len(steps) - 1:
            arrow_cell = table.rows[0].cells[2 * i + 1]
            arrow_cell.text = ""
            ap = arrow_cell.paragraphs[0]
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arun = ap.add_run("➜")
            arun.font.size = Pt(14)
            arun.font.color.rgb = ACCENT


# ── document body ──────────────────────────────────────────────────────

def build() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # ── Title block ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("Taqreerk")
    t_run.bold = True
    t_run.font.size = Pt(28)
    t_run.font.color.rgb = PRIMARY
    t_run.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run("RAG Architecture")
    s_run.font.size = Pt(18)
    s_run.font.color.rgb = ACCENT
    s_run.font.name = "Calibri"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m_run = meta.add_run(
        "Agentic RAG over PDF reports · LangGraph · Vertex AI · Docling · Langfuse · Ragas"
    )
    m_run.italic = True
    m_run.font.color.rgb = MUTED
    m_run.font.size = Pt(10)
    doc.add_paragraph()

    # ── 1. Overview ──
    add_heading(doc, "1. Overview", level=1)
    add_paragraph(
        doc,
        "Taqreerk answers questions about uploaded PDF reports with grounded, "
        "page-cited answers. The system is split across three services that share "
        "a single PostgreSQL+pgvector database:",
    )
    add_two_col_table(
        doc,
        ("Service", "Responsibility"),
        [
            (".NET 8 orchestrator (Tqreerk-backend)",
             "HTTP API, auth, RBAC, billing, quotas, file storage, "
             "ai_jobs lifecycle worker, persistence."),
            ("ai-service (Python / FastAPI)",
             "Agentic chat (LangGraph), retrieval, embedding, reranking, summarization, "
             "two-layer chat cache, Langfuse traces, Ragas evaluation."),
            ("doc-processor (Python / FastAPI, GPU)",
             "PDF extraction pipeline: layout, OCR, tables, figures, formulas, "
             "Arabic normalization, structure-aware chunking, embeddings, reranking."),
        ],
    )
    add_paragraph(
        doc,
        "The .NET layer never calls an LLM directly; it enqueues work to ai-service. "
        "ai-service in turn calls doc-processor for heavy extraction, then persists "
        "chunks + embeddings to Postgres so the agent can retrieve them at chat time.",
    )

    # ── 2. End-to-End Flow ──
    add_heading(doc, "2. End-to-End Flow", level=1)
    add_paragraph(doc, "Ingestion path:", bold=True)
    add_arrow_flow(doc, [
        "PDF upload",
        "doc-processor /v1/ingest_full",
        "Docling layout",
        "OCR + captions",
        "Chunk (struct-aware)",
        "Embed (Vertex)",
        "Persist report_chunks",
    ])
    doc.add_paragraph()
    add_paragraph(doc, "Chat path:", bold=True)
    add_arrow_flow(doc, [
        "User question",
        "ai-service /chat/{id}/messages (SSE)",
        "Cache lookup",
        "LangGraph agent",
        "Tool calls (≤5 hops)",
        "search_chunks → rerank",
        "Gemini generate",
        "Cited answer + Ragas score",
    ])

    # ── 3. LangGraph Agent ──
    add_heading(doc, "3. LangGraph Agentic Flow", level=1)
    add_paragraph(
        doc,
        "Defined in ai-service/pipelines/agent.py. The agent is a tool-using ReAct "
        "loop with hard cost guardrails — not a free-running planner. The graph has "
        "three nodes (START → agent → tools → agent → END) and a conditional edge "
        "that routes back to the agent after each tool batch or to END.",
    )
    add_two_col_table(
        doc,
        ("Element", "Detail"),
        [
            ("Graph topology",
             "Two real nodes — agent and tools — plus START/END. Agent calls Gemini; "
             "if the model emits tool_calls, edge routes to ToolNode; ToolNode runs "
             "all tools in parallel and routes back to agent."),
            ("State (AgentState)",
             "messages (LangChain reducer-managed history), hop_count (int), "
             "user_id, session_id. Auth identity for tools is closure-captured "
             "in ToolContext, not exposed to the LLM."),
            ("Hop cap",
             "MAX_TOOL_HOPS = 5. _route_after_agent ends the loop once the cap is "
             "hit even if the model still wants tools — prevents runaway loops "
             "and unbounded cost on a single turn."),
            ("Per-turn dedup",
             "A (tool_name, args_signature) cache short-circuits repeated identical "
             "tool calls within the same turn with a stable error message, so the "
             "model is forced to try a different angle instead of looping."),
            ("Fallback ladder",
             "Encoded in the 128-line system prompt: structured shortcuts "
             "(get_report_summary / get_report_keywords / get_report_indicators) "
             "→ semantic search_chunks → empty / unknown answer. Tool failures "
             "surface as ToolMessage(status=error) so the model can pivot."),
            ("Model",
             "gemini-2.5-flash via langchain-google-vertexai ChatVertexAI. "
             "temperature 0.2, max_output_tokens 2048."),
            ("Streaming",
             "astream_events(version=\"v2\") emits SSE events to the .NET caller: "
             "tool_call, tool_result, sources, token, done."),
        ],
    )

    # ── 4. Tools ──
    add_heading(doc, "4. Agent Tools", level=1)
    add_paragraph(
        doc,
        "Defined in ai-service/services/tools.py. 14 typed StructuredTools, each "
        "with a Pydantic schema. Auth identity (user_id, session_id) is wrapped in "
        "closure so the LLM can never spoof it. Every tool output is hard-capped at "
        "8000 characters with a “…[truncated; ask a narrower question]” marker.",
    )
    add_three_col_table(
        doc,
        ("Tool", "Purpose", "Notes"),
        [
            ("search_chunks",
             "Hybrid dense + BM25 retrieval over the report's chunks.",
             "RRF fusion of vector + tsvector hits, optional Vertex rerank, returns "
             "neighbor chunks for context. The workhorse retrieval tool."),
            ("get_page",
             "Fetch one page's full concatenated text.",
             "Direct SQL on report_chunks by page_number."),
            ("list_reports",
             "Filtered metadata search across all reports.",
             "Filters: sector, country, year, keyword. SQL joins."),
            ("get_report_metadata",
             "Title, org, sector, country, year, page count, ratings, status.",
             "Single report row + lookups."),
            ("get_report_summary",
             "AI-pre-generated summary + key findings (bilingual).",
             "Reads ReportAiContent.Summary."),
            ("get_report_indicators",
             "Extracted KPIs / numeric indicators (jsonb).",
             "ReportAiContent.Indicators."),
            ("get_report_trends",
             "Extracted trends with direction / time_span / magnitude.",
             "ReportAiContent.Trends."),
            ("get_report_recommendations",
             "Extracted action items / recommendations.",
             "ReportAiContent.Recommendations."),
            ("get_report_keywords",
             "Per-report auto-tagged keywords (per language).",
             "report_keywords table."),
            ("get_report_topics",
             "High-level topics / sectors per language.",
             "ReportAiContent.Topics."),
            ("get_translation",
             "Translated title / description / summary text — no download URL.",
             "Locked-in policy: never returns a file link from the agent."),
            ("list_saved_reports",
             "Caller's bookmarks (Published-only).",
             "Auth from ToolContext, never from LLM args."),
            ("list_user_interests",
             "Caller's followed sectors / orgs / countries.",
             "user_interests + LEFT JOINs."),
            ("find_similar_reports",
             "Embedding centroid cosine + topic / keyword overlap.",
             "score = cosine + 0.05 × shared_topics + 0.02 × shared_keywords."),
            ("get_session_history",
             "Recent turns in the current chat session.",
             "Used when the model needs to re-resolve a pronoun / follow-up."),
        ],
    )

    # ── 5. Extraction (doc-processor) ──
    add_heading(doc, "5. Extraction Pipeline (doc-processor)", level=1)
    add_paragraph(
        doc,
        "doc-processor/pipeline/orchestrator.py — process_page runs per rendered "
        "PNG and emits structured blocks the chunker can pack. The flow is "
        "layout-first: Docling identifies regions (text, heading, table, figure, "
        "formula, footer) with reading order, then each region dispatches to a "
        "specialist.",
    )
    add_three_col_table(
        doc,
        ("Stage", "Library / Model", "Notes"),
        [
            ("Layout + reading order + table structure",
             "Docling (DocLayNet + TableFormer)",
             "Built-in models. Source of region kind, bbox, text, reading_order, "
             "and table cell grids."),
            ("OCR",
             "EasyOCR (ar + en)",
             "Used as fallback when Docling text < ocr_fallback_min_chars (default "
             "20). Languages set via OCR_LANGUAGES env var."),
            ("Figure captioning",
             "Microsoft Florence-2 (base-ft)",
             "Task=<MORE_DETAILED_CAPTION>. Plus label OCR for axis text. Model id "
             "in doc-processor/core/config.py."),
            ("Formulas",
             "pix2tex (currently disabled)",
             "Formula regions are passed through; LaTeX recognition was tested "
             "but disabled in requirements.txt."),
            ("Tables → markdown",
             "Custom grid_to_markdown (tables.py)",
             "Renders TableFormer cells as GitHub-flavored markdown so the LLM can "
             "read row/col structure cheaply."),
            ("Arabic normalization",
             "arabic_normalize.py",
             "NFKC + visual→logical reorder for lines >60% Arabic. Fixes the "
             "PDF-extraction RTL issue without breaking mixed-direction lines."),
        ],
    )
    add_paragraph(
        doc,
        "Output of process_page: a list of blocks {type, content, reading_order, "
        "bbox} plus per-page metadata (section_title, page_type, language). This "
        "is what the chunker consumes.",
    )

    # ── 6. Chunking ──
    add_heading(doc, "6. Chunking Method", level=1)
    add_paragraph(
        doc,
        "Two sibling files — ai-service/core/chunking.py and "
        "doc-processor/pipeline/chunker.py — kept byte-identical and verified by "
        "scripts/check_chunker_sync.py in CI. The same algorithm is used whether "
        "extraction was Gemini-Vision-only (chunk_text) or doc-processor "
        "(chunk_blocks_with_meta).",
    )
    add_paragraph(doc, "Algorithm — chunk_blocks_with_meta:", bold=True)
    add_bullet(doc, "Sort blocks by reading_order so chunks reflect the page's natural reading flow.")
    add_bullet(doc, "Atomic blocks: tables, figures, and formulas are NEVER split — each is emitted as its own chunk regardless of size, so structure is preserved for retrieval.")
    add_bullet(doc, "Headings: start a new chunk and become the section_title for subsequent chunks until the next heading replaces them.")
    add_bullet(doc, "Text blocks: greedy-pack into the current chunk; if a single text block exceeds the limit, RecursiveCharacterTextSplitter splits it on paragraph → line → sentence (Arabic-aware separators ؟ ۔ included) → word → char.")
    add_bullet(doc, "Each emitted chunk carries content + block_types[] + section_title — block_types let search_chunks filter (e.g. block_types=[\"table\"] for KPI questions).")
    add_two_col_table(
        doc,
        ("Parameter", "Value"),
        [
            ("DEFAULT_CHUNK_CHARS", "2000 (~500 tokens for mixed AR+EN)"),
            ("CHUNK_OVERLAP", "200 (~50 tokens)"),
            ("Char-to-token ratio (AR+EN mix)", "≈ 3–4 chars / token (empirical)"),
            ("Splitter", "LangChain RecursiveCharacterTextSplitter"),
            ("Sync rule", "byte-identical with doc-processor/pipeline/chunker.py — "
                          "enforced by scripts/check_chunker_sync.py (CI-gated)."),
        ],
    )

    # ── 7. Embedding & Reranking ──
    add_heading(doc, "7. Embedding & Reranking", level=1)
    add_paragraph(doc, "Embedding (production):", bold=True)
    add_two_col_table(
        doc,
        ("Aspect", "Value"),
        [
            ("Model", "Vertex AI gemini-embedding-001 (managed, always warm)."),
            ("Dimensions", "768 (Matryoshka-truncated from native dim)."),
            ("Task type",
             "RETRIEVAL_DOCUMENT for chunks, RETRIEVAL_QUERY for the user question. "
             "These are NOT interchangeable — the spaces are tuned asymmetrically."),
            ("Batch size", "≤ 250 texts per Vertex call (services/embed.py)."),
            ("Why not BGE-M3 in prod?",
             "BGE-M3 (still present in doc-processor) had strong Arabic/English "
             "MIRACL but cold-start GPU latency (30–90 s) broke first-query UX. "
             "Vertex is sub-150 ms typical and the dim-768 cost fits the budget."),
            ("Re-ingest policy",
             "Embedding-model changes truncate report_chunks AND chat_cache. "
             "Vector spaces are not portable — see migration "
             "20260430210000_VertexEmbeddingDimensions.cs."),
        ],
    )
    add_paragraph(doc, "Reranking (active):", bold=True)
    add_two_col_table(
        doc,
        ("Aspect", "Value"),
        [
            ("Model (ai-service)",
             "Vertex AI Ranking API — semantic-ranker-default-004."),
            ("Pool / top-k", "Fetch 20 candidates, rerank to top 5."),
            ("Failure mode",
             "Soft fail — on rerank API error, original retrieval order is "
             "returned with score=0 and the chat continues."),
            ("Fallback model (doc-processor)",
             "BAAI/bge-reranker-v2-m3 cross-encoder on GPU — used by the "
             "doc-processor /v1/rerank endpoint. Sibling to bge-m3."),
        ],
    )

    # ── 8. Retrieval ──
    add_heading(doc, "8. Retrieval (search_chunks internals)", level=1)
    add_paragraph(
        doc,
        "search_chunks combines three signals before reranking, all scoped to a "
        "single ReportId. Cross-report bleed is structurally impossible — it's a "
        "WHERE clause in every query, not a runtime filter the agent can omit.",
    )
    add_bullet(doc, "Dense: cosine similarity on the 768-dim embedding column (pgvector).")
    add_bullet(doc, "Lexical: PostgreSQL tsvector full-text using both 'simple' and 'arabic' configurations (bilingual).")
    add_bullet(doc, "Metadata filter: optional jsonb filter (e.g. block_types contains 'table') accelerated by a GIN index with jsonb_path_ops.")
    add_bullet(doc, "Fusion: Reciprocal Rank Fusion (RRF) merges the two ranked lists into a single candidate pool.")
    add_bullet(doc, "Rerank: top 20 → Vertex Ranking API → top 5.")
    add_bullet(doc, "Neighbor expansion: each surviving chunk is returned with adjacent chunks on the same page so the LLM gets continuity, not isolated sentences.")

    add_paragraph(doc, "Indexes on report_chunks:", bold=True)
    add_two_col_table(
        doc,
        ("Index", "Purpose"),
        [
            ("IX_report_chunks_ReportId_PageNumber_ChunkIndex (unique)",
             "Deduplication — one chunk per (report, page, chunk index)."),
            ("IX_report_chunks_ReportId_PageNumber",
             "Page-level retrieval (e.g. get_page tool)."),
            ("IX_report_chunks_search_vector (GIN)",
             "Bilingual tsvector full-text lookup, sub-millisecond."),
            ("idx_report_chunks_metadata_gin (jsonb_path_ops, GIN)",
             "Metadata-filtered queries — added concurrently to avoid table lock."),
            ("Vector ANN index (pgvector, cosine ops)",
             "Approximate nearest-neighbor on the 768-dim embedding column."),
        ],
    )

    # ── 9. Chat Cache ──
    add_heading(doc, "9. Two-Layer Chat Cache", level=1)
    add_paragraph(
        doc,
        "ai-service/services/chat_cache.py. Owned by ai-service; the .NET layer "
        "never reads or writes it. Both layers short-circuit retrieval AND LLM "
        "generation.",
    )
    add_three_col_table(
        doc,
        ("Layer", "Mechanism", "TTL / threshold"),
        [
            ("L1 — Exact",
             "PRIMARY KEY = SHA-256(report_id ‖ normalized_question). One indexed "
             "select; sub-millisecond hit.",
             "exact_ttl = 86 400 s (24 h)"),
            ("L2 — Semantic",
             "Cosine similarity on cached question_emb (vector 768). Only runs "
             "after L1 miss AND only when the caller is already embedding the "
             "question (no extra cost).",
             "threshold ≥ 0.95 cosine; semantic_ttl = 3 600 s (1 h)"),
        ],
    )
    add_bullet(doc, "Schema: cache_key, report_id, question, question_emb, answer, source_pages (jsonb), hit_count, expires_at.")
    add_bullet(doc, "Normalization is conservative: lowercase, strip trailing punctuation (Arabic ؟؛ + English), collapse whitespace — avoids false hits on subtly different questions.")
    add_bullet(doc, "hit_count gives a popularity signal for content/UX feedback.")

    # ── 10. Langfuse Observability ──
    add_heading(doc, "10. Observability (Langfuse)", level=1)
    add_paragraph(
        doc,
        "ai-service/services/observability.py. A lazy-init singleton wraps the "
        "Langfuse client. If Langfuse is disabled or misconfigured, every helper "
        "returns a no-op handle so a Langfuse outage cannot break a user request.",
    )
    add_two_col_table(
        doc,
        ("Primitive", "Use"),
        [
            ("trace()",
             "Starts a request-scoped trace with user_id, session_id, tags "
             "(e.g. [chat, agent]); honors langfuse_trace_sample_rate."),
            ("span()",
             "Context-managed span for inner stages (retrieval, rerank, cache "
             "lookup); auto-ends and marks errors."),
            ("generation()",
             "Logs every LLM call — model id, prompt, completion, token usage. "
             "Wired into the LangChain CallbackHandler attached to the graph, so "
             "every Gemini call inside the agent loop is captured."),
            ("score()",
             "Posts numeric scores (Ragas faithfulness / answer_relevancy / "
             "context_precision) to the existing trace."),
        ],
    )
    add_bullet(doc, "Sampling: per-call Bernoulli on langfuse_trace_sample_rate (default 1.0 in staging).")
    add_bullet(doc, "Flush: batched every 5 s or 100 events; explicit flush at SSE close; atexit handler to drain on process exit.")
    add_bullet(doc, "Defense-in-depth: every helper has a no-op fallback so observability failures stay silent for the user.")

    # ── 11. Ragas Evaluation ──
    add_heading(doc, "11. Online Evaluation (Ragas)", level=1)
    add_paragraph(
        doc,
        "ai-service/pipelines/eval.py and ai-service/services/eval_models.py. "
        "Evaluation is reference-free and runs on live traffic — every chat turn "
        "(subject to eval_sample_rate) is scored asynchronously after the SSE "
        "stream closes, so the user never waits for it.",
    )
    add_two_col_table(
        doc,
        ("Metric", "What it measures"),
        [
            ("faithfulness",
             "Does the answer derive from the retrieved contexts? Catches "
             "hallucinations."),
            ("answer_relevancy",
             "Does the answer actually address the question (vs. tangential)?"),
            ("context_precision_without_reference",
             "Are the retrieved chunks on-topic? Catches retrieval drift."),
        ],
    )
    add_two_col_table(
        doc,
        ("Decision", "Detail"),
        [
            ("Judge LLM",
             "Gemini Flash via a custom GeminiRagasLLM wrapper — reuses the "
             "ai-service retry stack instead of Ragas's bare client."),
            ("Judge embeddings",
             "Same gemini-embedding-001 (via DocProcessorRagasEmbeddings) so "
             "judge and production share the vector space."),
            ("Dataset",
             "Live traces — questions / answers from chat_messages, contexts "
             "from the tool results captured in the trace. No golden set."),
            ("Triggering",
             "Enqueued by pipelines/jobs.py after the SSE finishes; worker "
             "calls run_eval(); fully out-of-band from the user request."),
            ("Surface",
             "Posted as Langfuse scores on the same trace_id. Stable score "
             "names (faithfulness / answer_relevancy / context_precision) so "
             "dashboards survive Ragas version bumps. No DB write."),
            ("Timeout protection",
             "Three-layer: per-metric timeout (default 90 s), Ragas internal "
             "max_retries=2, and an outer asyncio.wait_for() hard cap. "
             "Prevents a stuck judge call from blocking the worker pool."),
        ],
    )

    # ── 12. Prompts ──
    add_heading(doc, "12. Prompts", level=1)
    add_paragraph(
        doc,
        "ai-service/core/prompts.py + the inline SYSTEM_PROMPT in "
        "pipelines/agent.py.",
    )
    add_two_col_table(
        doc,
        ("Prompt", "Purpose"),
        [
            ("SYSTEM_PROMPT (agent)",
             "128 lines: tool-calling discipline, fallback ladder, multi-turn "
             "reformulation rules, search-query selection, citation guidance, "
             "max-5-hops cost rule."),
            ("PAGE_DESCRIPTION",
             "Vision prompt: extract text + visual elements + section title + "
             "page type + language from a rendered page (Gemini Vision path)."),
            ("chat_system_prompt",
             "Per-turn system prompt that embeds the retrieved context block "
             "before the user's message in the answer LLM call."),
            ("summarize_prompt",
             "One-shot Gemini call returning summary + key_findings + topics + "
             "indicators + trends — bound by REPORT_SUMMARY_SCHEMA."),
            ("insights_prompt",
             "Lighter call: indicators + trends only. Bound by INSIGHTS_SCHEMA."),
            ("compare_prompt",
             "Multi-report comparison. Bound by COMPARE_SCHEMA."),
        ],
    )

    # ── 13. Service Endpoints ──
    add_heading(doc, "13. Service Endpoints", level=1)
    add_paragraph(doc, "ai-service (consumed by .NET and the SPA chat):", bold=True)
    add_two_col_table(
        doc,
        ("Endpoint", "Purpose"),
        [
            ("POST /chat/sessions",
             "Create chat session for (user_id, report_id). Returns session_id "
             "and an initial summary."),
            ("POST /chat/{session_id}/messages",
             "Send a question. Streams SSE events: tool_call, tool_result, "
             "sources, token, done."),
            ("GET  /chat/{session_id}/history",
             "Recent turns for the session."),
            ("POST /reports/ingest",
             "Fire-and-forget (HTTP 202). Returns job_id."),
            ("POST /reports/translate",
             "Synchronous translation (timeout 300 s)."),
            ("GET  /reports/jobs/{job_id}",
             "Polled by .NET ReportProcessingWorker every 5 s."),
        ],
    )
    add_paragraph(doc, "doc-processor (consumed by ai-service):", bold=True)
    add_two_col_table(
        doc,
        ("Endpoint", "Purpose"),
        [
            ("POST /v1/extract",
             "One page PNG → blocks + content + metadata."),
            ("POST /v1/extract_document",
             "Full PDF → per-page extracted blocks."),
            ("POST /v1/ingest_full",
             "Full PDF → chunks + embeddings, ready to persist. "
             "Default ingest path in production."),
            ("POST /v1/embed",
             "BGE-M3 embedding endpoint (used as fallback / for parity tests)."),
            ("POST /v1/rerank",
             "BGE-reranker-v2-m3 cross-encoder rerank."),
            ("GET  /health", "Model readiness."),
        ],
    )
    add_paragraph(
        doc,
        "All doc-processor routes validate X-Internal-Token (defense-in-depth) "
        "when settings.internal_api_token is set.",
        italic=True,
    )

    # ── 14. Models & Config ──
    add_heading(doc, "14. Models & Configuration", level=1)
    add_paragraph(doc, "ai-service (core/config.py):", bold=True)
    add_two_col_table(
        doc,
        ("Setting", "Default"),
        [
            ("gemini_chat_model", "gemini-2.5-flash"),
            ("gemini_vision_model", "gemini-2.5-flash"),
            ("gemini_summary_model", "gemini-2.5-flash"),
            ("gemini_embed_model", "gemini-embedding-001"),
            ("gemini_embed_dim", "768"),
            ("reranker_vertex_model", "semantic-ranker-default-004"),
            ("reranker_enabled / candidate_pool / top_k", "True / 20 / 5"),
            ("langfuse_enabled / trace_sample_rate", "True / 1.0"),
            ("eval_enabled / sample_rate / metric_timeout_s", "True / 1.0 / 90"),
            ("chat_cache_enabled / exact_ttl / semantic_threshold",
             "True / 86400 s / 0.95"),
            ("doc_processor_enabled / ingest_full_enabled",
             "False / True (i.e. when doc-processor is on, use /v1/ingest_full)"),
        ],
    )
    add_paragraph(doc, "doc-processor (core/config.py):", bold=True)
    add_two_col_table(
        doc,
        ("Setting", "Default"),
        [
            ("florence_model_id", "microsoft/Florence-2-base-ft"),
            ("ocr_languages", "ar,en"),
            ("embed_model_id", "BAAI/bge-m3"),
            ("rerank_model_id", "BAAI/bge-reranker-v2-m3"),
            ("embed_max_seq_length / batch", "8192 / 16"),
            ("rerank_max_seq_length / batch", "1024 / 8"),
            ("device / fp16", "cuda / True"),
            ("ocr_fallback_min_chars", "20"),
            ("max_blocks_per_page", "200"),
            ("embed_vertex_model / dim",
             "gemini-embedding-001 / 768 (used by /v1/ingest_full bridge)"),
        ],
    )
    add_paragraph(doc, "Environment variables shared across services:", bold=True)
    add_code_block(
        doc,
        "DATABASE_URL                  # shared by .NET, ai-service, doc-processor\n"
        "GCP_PROJECT_ID, VERTEX_LOCATION\n"
        "GEMINI_API_KEY                # optional (AI Studio fallback)\n"
        "LANGFUSE_HOST / PUBLIC_KEY / SECRET_KEY\n"
        "SENTRY_DSN                    # error reporting\n"
        "INTERNAL_API_TOKEN            # doc-processor route guard",
    )

    # ── 15. Operational Notes ──
    add_heading(doc, "15. Operational Notes", level=1)
    add_bullet(doc, "Idempotency: ingest is keyed by report_id; re-runs supersede prior chunks for that report.")
    add_bullet(doc, "Quotas (.NET-side): DailyIngestPerOrg=20, DailyTranslatePerOrg=10, DailyChatPerUser=200 — gate enqueues before HTTP 202.")
    add_bullet(doc, "Cost guardrails (ai-service): MAX_TOOL_HOPS=5 caps per-turn tool usage; per-turn dedup cache prevents repeated identical tool calls.")
    add_bullet(doc, "Failure-soft: chat cache, Langfuse, and reranking each fail open — production never breaks because of an observability or cache outage.")
    add_bullet(doc, "Eval is async post-stream — never on the user's critical path; metric timeouts ensure a stuck judge can't starve the worker pool.")
    add_bullet(doc, "Embedding-model swaps require a full re-ingest (vector spaces are not portable) — the migration explicitly truncates report_chunks and chat_cache.")
    add_bullet(doc, "Chunker drift between ai-service and doc-processor is caught at PR time by scripts/check_chunker_sync.py; the two files are byte-identical.")
    add_bullet(doc, "All retrieval is report-scoped at the SQL level — multi-tenant isolation is structural, not a runtime filter the agent can omit.")

    # ── 16. Key Source Files ──
    add_heading(doc, "16. Key Source Files", level=1)
    add_two_col_table(
        doc,
        ("Path", "Role"),
        [
            ("ai-service/pipelines/agent.py",
             "LangGraph definition, AgentState, SYSTEM_PROMPT, hop cap, dedup cache."),
            ("ai-service/services/tools.py",
             "All 14 StructuredTools, ToolContext, output cap, auth wrapping."),
            ("ai-service/services/observability.py",
             "Langfuse singleton, trace/span/generation/score helpers, no-op fallback."),
            ("ai-service/pipelines/eval.py + services/eval_models.py",
             "Ragas runner, GeminiRagasLLM, DocProcessorRagasEmbeddings, "
             "timeout layers."),
            ("ai-service/core/chunking.py", "Production chunker (sibling of doc-processor)."),
            ("ai-service/services/embed.py / reranker.py",
             "Vertex embedding client, Vertex Ranking API client."),
            ("ai-service/services/chat_cache.py",
             "Two-layer (exact SHA-256 + semantic cosine) chat cache."),
            ("ai-service/core/prompts.py", "Vision, summary, insights, compare prompts + JSON schemas."),
            ("ai-service/api/chat.py", "SSE chat endpoint, ToolContext assembly, stream events."),
            ("doc-processor/pipeline/orchestrator.py",
             "Per-page extraction loop (Docling → OCR / tables / figures / formulas)."),
            ("doc-processor/pipeline/{ocr,tables,figures,formulas,layout,arabic_normalize}.py",
             "Specialist stages."),
            ("doc-processor/pipeline/chunker.py",
             "Byte-identical sibling of ai-service chunker."),
            ("doc-processor/pipeline/{embeddings,vertex_embedder,reranker}.py",
             "BGE-M3 / Vertex embed bridge / BGE-reranker-v2-m3."),
            ("doc-processor/api/extract.py + main.py",
             "FastAPI routes: /v1/extract, /v1/ingest_full, /v1/embed, /v1/rerank."),
            ("scripts/check_chunker_sync.py",
             "CI guard — fails the build if the two chunkers drift."),
            ("Tqreerk-backend/Application/Services/ReportAiService.cs + "
             "Infrastructure/AI/AiServiceClient.cs",
             ".NET-side enqueue + typed HttpClient to ai-service."),
            ("Tqreerk-backend/Application/Workers/ReportProcessingWorker.cs",
             "Polls ai-service jobs, advances Report.Status."),
        ],
    )

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
